"""
extract.py — Step 2-4: pipeline di estrazione documenti .docx → SQLite

Uso:
    python extract.py                        # processa tutti i file in rules.yaml
    python extract.py --only fontana         # solo i file il cui label contiene "fontana"
    python extract.py --only "MES Web Servizi"
    python extract.py --dry-run              # simula senza scrivere sul DB

Output:
    data/mes_docs.db  — database SQLite con tutto il contenuto indicizzato
"""

import argparse
import hashlib
import io
import logging
import re
import sqlite3
import sys
import time
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path

import yaml

try:
    import docx
except ImportError:
    print("Errore: pip install python-docx pyyaml")
    sys.exit(1)

# ---------------------------------------------------------------------------
# Schema SQL
# ---------------------------------------------------------------------------

SCHEMA = """
CREATE TABLE IF NOT EXISTS documents (
    id              INTEGER PRIMARY KEY AUTOINCREMENT,
    filename        TEXT NOT NULL,
    label           TEXT,
    full_path       TEXT NOT NULL,
    revision        TEXT,
    tags            TEXT,           -- JSON array come stringa
    extracted_at    TEXT NOT NULL,
    file_hash       TEXT,           -- MD5 del file per incremental indexing
    paragraphs_n    INTEGER DEFAULT 0,
    tables_n        INTEGER DEFAULT 0,
    parameters_n    INTEGER DEFAULT 0
);

CREATE TABLE IF NOT EXISTS sections (
    id              INTEGER PRIMARY KEY AUTOINCREMENT,
    document_id     INTEGER NOT NULL REFERENCES documents(id),
    parent_id       INTEGER REFERENCES sections(id),
    level           INTEGER NOT NULL,   -- 1=H1, 2=H2, ecc.
    title           TEXT,
    order_index     INTEGER NOT NULL
);

CREATE TABLE IF NOT EXISTS paragraphs (
    id              INTEGER PRIMARY KEY AUTOINCREMENT,
    document_id     INTEGER NOT NULL REFERENCES documents(id),
    section_id      INTEGER REFERENCES sections(id),
    style           TEXT,
    text            TEXT NOT NULL,
    order_index     INTEGER NOT NULL,
    char_count      INTEGER
);

CREATE TABLE IF NOT EXISTS parameters (
    id              INTEGER PRIMARY KEY AUTOINCREMENT,
    document_id     INTEGER NOT NULL REFERENCES documents(id),
    section_id      INTEGER REFERENCES sections(id),
    name            TEXT NOT NULL,
    value           TEXT,
    pattern_name    TEXT,
    confidence      TEXT,           -- HIGH / MEDIUM / LOW
    raw_text        TEXT,
    order_index     INTEGER NOT NULL
);

CREATE TABLE IF NOT EXISTS table_cells (
    id              INTEGER PRIMARY KEY AUTOINCREMENT,
    document_id     INTEGER NOT NULL REFERENCES documents(id),
    section_id      INTEGER REFERENCES sections(id),
    table_index     INTEGER NOT NULL,
    row_index       INTEGER NOT NULL,
    col_index       INTEGER NOT NULL,
    cell_text       TEXT
);

-- Full-text search: tabella standalone popolata esplicitamente
-- Colonne UNINDEXED non vengono indicizzate ma sono disponibili nei risultati
CREATE VIRTUAL TABLE IF NOT EXISTS fts_paragraphs USING fts5(
    text,
    doc_label    UNINDEXED,
    doc_filename UNINDEXED,
    section_title UNINDEXED,
    paragraph_id  UNINDEXED,
    tokenize = 'unicode61 remove_diacritics 1'
);
"""

# ---------------------------------------------------------------------------
# Pattern rilevamento parametri (specchio di rules.yaml per comodità)
# ---------------------------------------------------------------------------

@dataclass
class ParamPattern:
    name: str
    regex: re.Pattern | None
    match_type: str   # "line" | "inline" | "table"

PARAM_PATTERNS: list[ParamPattern] = [
    ParamPattern("shell_export",
                 re.compile(r'^export\s+([A-Z][A-Z0-9_]{2,})\s*=\s*(.+)$'),
                 "line"),
    ParamPattern("key_value_uppercase",
                 re.compile(r'^([A-Z][A-Z0-9_]{2,})\s*=\s*(.+)$'),
                 "line"),
    ParamPattern("key_value_snakecase",
                 re.compile(r'^([a-z][a-z0-9_]{2,})\s*=\s*([0-9]+|\'[^\']*\'|"[^"]*"|[A-Z])\s*$'),
                 "line"),
    ParamPattern("inline_uppercase_param",
                 re.compile(r'([A-Z][A-Z0-9_]{2,})\s*=\s*([0-9]+)'),
                 "inline"),
]

NARRATIVE_PREFIXES = re.compile(
    r'^\s*(se |quando |qualora |esempio:|ad esempio|nel caso|if |nota:|nbb |nbb:|attenzione)',
    re.IGNORECASE
)
NARRATIVE_OPERATORS = re.compile(
    r'\b(è|sono|deve|vale|vuol dire|significa|ovvero|cioè|corrisponde)\b',
    re.IGNORECASE
)
SQL_CONTEXT = re.compile(
    r'\b(SELECT|FROM|WHERE|AND|OR|INSERT|UPDATE|DELETE|JOIN)\b',
    re.IGNORECASE
)


def compute_confidence(text: str, pattern_name: str, key: str, value: str) -> str:
    score = 0
    if "_" in key:
        score += 1
    if re.match(r'^[0-9]+$', value):
        score += 1
    if len(text.strip()) < 80:
        score += 1
    if pattern_name in ("shell_export", "key_value_uppercase", "key_value_snakecase"):
        score += 1
    if NARRATIVE_PREFIXES.match(text.strip()):
        score -= 2
    if NARRATIVE_OPERATORS.search(text.strip()):
        score -= 1
    if SQL_CONTEXT.search(text.strip()):
        score -= 3   # quasi certamente falso positivo
    if pattern_name == "inline_uppercase_param" and len(text.strip()) > 100:
        score -= 1
    if score >= 3:
        return "HIGH"
    if score >= 1:
        return "MEDIUM"
    return "LOW"


@dataclass
class DetectedParam:
    name: str
    value: str
    pattern_name: str
    confidence: str
    raw_text: str


def detect_params(text: str) -> list[DetectedParam]:
    text_s = text.strip()
    results: list[DetectedParam] = []
    seen: set[str] = set()

    for pp in PARAM_PATTERNS:
        if pp.match_type == "line":
            m = pp.regex.match(text_s)
            if m:
                key, val = m.group(1), m.group(2).strip()
                if key not in seen:
                    seen.add(key)
                    conf = compute_confidence(text_s, pp.name, key, val)
                    results.append(DetectedParam(key, val, pp.name, conf, text_s))
                break
        else:  # inline
            for m in pp.regex.finditer(text_s):
                key, val = m.group(1), m.group(2).strip()
                if key not in seen:
                    seen.add(key)
                    conf = compute_confidence(text_s, pp.name, key, val)
                    results.append(DetectedParam(key, val, pp.name, conf, text_s))

    return results


def detect_table_params(table, section_id: int, document_id: int,
                        table_index: int) -> list[DetectedParam]:
    """Parametri in tabelle 2-colonne: cella[0]=nome, cella[1]=valore."""
    results = []
    for row in table.rows:
        if len(row.cells) >= 2:
            c0 = row.cells[0].text.strip()
            c1 = row.cells[1].text.strip()
            if re.match(r'^[A-Z][A-Z0-9_]{2,}$', c0) and c1:
                conf = compute_confidence(f"{c0}={c1}", "table_key_value", c0, c1)
                results.append(DetectedParam(c0, c1, "table_key_value", conf, f"{c0} | {c1}"))
    return results


# ---------------------------------------------------------------------------
# Sezioni gerarchiche
# ---------------------------------------------------------------------------

HEADING_STYLES = {
    "Heading 1": 1, "Heading 2": 2, "Heading 3": 3,
    "Heading 4": 4, "Heading 5": 5,
    "Titolo 1": 1, "Titolo 2": 2, "Titolo 3": 3,
}


@dataclass
class SectionStack:
    """Tiene traccia della gerarchia di sezioni corrente durante la lettura."""
    stack: list[tuple[int, int]] = field(default_factory=list)
    # ogni elemento: (level, section_db_id)

    def push(self, level: int, section_id: int):
        # rimuove tutto quello che è allo stesso livello o più profondo
        self.stack = [(l, sid) for l, sid in self.stack if l < level]
        self.stack.append((level, section_id))

    def current_id(self) -> int | None:
        return self.stack[-1][1] if self.stack else None

    def parent_id_for(self, level: int) -> int | None:
        parents = [(l, sid) for l, sid in self.stack if l < level]
        return parents[-1][1] if parents else None


# ---------------------------------------------------------------------------
# Extraction da un documento
# ---------------------------------------------------------------------------

def file_hash(path: str) -> str:
    h = hashlib.md5()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


def extract_document(conn: sqlite3.Connection, source: dict,
                     base_path: str, dry_run: bool,
                     log: logging.Logger) -> dict:
    """
    Estrae un documento e lo salva nel DB.
    Ritorna statistiche: {paragraphs, tables, parameters, duration_ms}.
    """
    full_path = str(Path(base_path) / source["path"])
    filename = Path(source["path"]).name
    label = source.get("label", filename)
    tags = str(source.get("tags", []))

    # Estrai revisione dal nome file (es. .r01. → "r01")
    rev_match = re.search(r'\.(r\d+)\.', filename, re.IGNORECASE)
    revision = rev_match.group(1) if rev_match else None

    log.info(f"Apertura: {filename}")
    t0 = time.monotonic()

    try:
        doc_obj = docx.Document(full_path)
    except Exception as e:
        log.error(f"  Impossibile aprire {filename}: {e}")
        return {}

    fhash = file_hash(full_path)
    extracted_at = datetime.now().isoformat()

    if dry_run:
        log.info(f"  [dry-run] {filename} — {len(doc_obj.paragraphs)} paragrafi, {len(doc_obj.tables)} tabelle")
        return {"paragraphs": len(doc_obj.paragraphs), "tables": len(doc_obj.tables),
                "parameters": 0, "duration_ms": 0}

    cur = conn.cursor()

    # --- Inserisci documento ---
    cur.execute("""
        INSERT INTO documents (filename, label, full_path, revision, tags, extracted_at, file_hash)
        VALUES (?, ?, ?, ?, ?, ?, ?)
    """, (filename, label, full_path, revision, tags, extracted_at, fhash))
    doc_id = cur.lastrowid

    section_stack = SectionStack()
    para_count = 0
    param_count = 0
    section_order = 0

    # --- Itera paragrafi ---
    for order_idx, p in enumerate(doc_obj.paragraphs):
        text = p.text.strip()
        style = p.style.name

        # Heading → nuova sezione
        if style in HEADING_STYLES:
            level = HEADING_STYLES[style]
            parent_id = section_stack.parent_id_for(level)
            cur.execute("""
                INSERT INTO sections (document_id, parent_id, level, title, order_index)
                VALUES (?, ?, ?, ?, ?)
            """, (doc_id, parent_id, level, text, section_order))
            sec_id = cur.lastrowid
            section_stack.push(level, sec_id)
            section_order += 1
            continue

        if not text:
            continue

        # Paragrafo normale
        cur.execute("""
            INSERT INTO paragraphs (document_id, section_id, style, text, order_index, char_count)
            VALUES (?, ?, ?, ?, ?, ?)
        """, (doc_id, section_stack.current_id(), style, text, order_idx, len(text)))
        para_count += 1

        # Rilevamento parametri nel testo
        for pm in detect_params(text):
            cur.execute("""
                INSERT INTO parameters
                    (document_id, section_id, name, value, pattern_name, confidence, raw_text, order_index)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """, (doc_id, section_stack.current_id(),
                  pm.name, pm.value, pm.pattern_name, pm.confidence, pm.raw_text, order_idx))
            param_count += 1

    # --- Itera tabelle ---
    for t_idx, table in enumerate(doc_obj.tables):
        # Determina la sezione corrente al momento della tabella
        # (le tabelle non hanno un indice lineare nei paragrafi,
        #  usiamo la sezione più recente trovata)
        current_sec = section_stack.current_id()

        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                if not cell_text:
                    continue
                cur.execute("""
                    INSERT INTO table_cells
                        (document_id, section_id, table_index, row_index, col_index, cell_text)
                    VALUES (?, ?, ?, ?, ?, ?)
                """, (doc_id, current_sec, t_idx, r_idx, c_idx, cell_text))

        # Parametri tabellari (col0=nome, col1=valore)
        for pm in detect_table_params(table, current_sec, doc_id, t_idx):
            cur.execute("""
                INSERT INTO parameters
                    (document_id, section_id, name, value, pattern_name, confidence, raw_text, order_index)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """, (doc_id, current_sec,
                  pm.name, pm.value, pm.pattern_name, pm.confidence, pm.raw_text, t_idx))
            param_count += 1

    # --- Aggiorna contatori sul documento ---
    cur.execute("""
        UPDATE documents
        SET paragraphs_n=?, tables_n=?, parameters_n=?
        WHERE id=?
    """, (para_count, len(doc_obj.tables), param_count, doc_id))

    conn.commit()

    duration_ms = int((time.monotonic() - t0) * 1000)
    log.info(f"  OK  paragrafi={para_count}  tabelle={len(doc_obj.tables)}  "
             f"parametri={param_count}  {duration_ms}ms")

    return {"paragraphs": para_count, "tables": len(doc_obj.tables),
            "parameters": param_count, "duration_ms": duration_ms}


# ---------------------------------------------------------------------------
# FTS5 index build
# ---------------------------------------------------------------------------

def build_fts(conn: sqlite3.Connection, log: logging.Logger):
    """Popola (o ricostruisce) l'indice FTS5 da paragraphs + metadati."""
    log.info("Costruzione indice FTS5...")
    cur = conn.cursor()
    cur.execute("DELETE FROM fts_paragraphs")
    cur.execute("""
        INSERT INTO fts_paragraphs (text, doc_label, doc_filename, section_title, paragraph_id)
        SELECT
            p.text,
            d.label,
            d.filename,
            COALESCE(s.title, ''),
            p.id
        FROM paragraphs p
        JOIN documents d ON p.document_id = d.id
        LEFT JOIN sections s ON p.section_id = s.id
        WHERE length(p.text) > 2
    """)
    n = cur.execute("SELECT count(*) FROM fts_paragraphs").fetchone()[0]
    conn.commit()
    log.info(f"FTS5 pronto: {n} righe indicizzate")


# ---------------------------------------------------------------------------
# Entrypoint
# ---------------------------------------------------------------------------

def setup_db(db_path: str, recreate: bool) -> sqlite3.Connection:
    Path(db_path).parent.mkdir(parents=True, exist_ok=True)
    if recreate and Path(db_path).exists():
        Path(db_path).unlink()
    conn = sqlite3.connect(db_path)
    conn.executescript(SCHEMA)
    conn.commit()
    return conn


def setup_logging(log_file: str, level: str) -> logging.Logger:
    Path(log_file).parent.mkdir(parents=True, exist_ok=True)
    log = logging.getLogger("extract")
    log.setLevel(getattr(logging, level.upper(), logging.INFO))
    fmt = logging.Formatter("%(asctime)s  %(levelname)-7s  %(message)s",
                            datefmt="%H:%M:%S")
    # console
    ch = logging.StreamHandler(
        io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace"))
    ch.setFormatter(fmt)
    log.addHandler(ch)
    # file
    fh = logging.FileHandler(log_file, encoding="utf-8")
    fh.setFormatter(fmt)
    log.addHandler(fh)
    return log


def main():
    parser = argparse.ArgumentParser(description="Estrai documenti .docx in SQLite")
    parser.add_argument("--config", default="rules.yaml",
                        help="Path al file rules.yaml (default: rules.yaml)")
    parser.add_argument("--only",  default=None,
                        help="Processa solo i file il cui label contiene questa stringa")
    parser.add_argument("--dry-run", action="store_true",
                        help="Simula l'estrazione senza scrivere sul DB")
    parser.add_argument("--build-fts", action="store_true",
                        help="(Ri)costruisce solo l'indice FTS5 senza riestrarre i documenti")
    args = parser.parse_args()

    cfg = yaml.safe_load(open(args.config, encoding="utf-8"))
    log = setup_logging(cfg["logging"]["log_file"], cfg["logging"]["level"])

    log.info("=" * 60)
    log.info("Avvio estrazione")
    log.info(f"Config: {args.config}")
    log.info(f"DB: {cfg['database']['path']}")
    log.info(f"Dry-run: {args.dry_run}")

    conn = None
    if not args.dry_run:
        conn = setup_db(cfg["database"]["path"], cfg["database"]["recreate_on_run"])
        log.info("Schema DB creato/verificato")

    # Modalità solo FTS rebuild
    if args.build_fts:
        if conn:
            build_fts(conn, log)
            conn.close()
        return

    sources = cfg["sources"]
    if args.only:
        sources = [s for s in sources
                   if args.only.lower() in s.get("label", "").lower()
                   or args.only.lower() in s["path"].lower()]
        log.info(f"Filtro '--only {args.only}': {len(sources)} file selezionati")

    total = {"paragraphs": 0, "tables": 0, "parameters": 0, "duration_ms": 0, "files": 0}

    for source in sources:
        stats = extract_document(conn, source, cfg["base_path"], args.dry_run, log)
        if stats:
            for k in ("paragraphs", "tables", "parameters", "duration_ms"):
                total[k] += stats.get(k, 0)
            total["files"] += 1

    log.info("-" * 60)
    log.info(f"Completato: {total['files']} file  |  "
             f"paragrafi={total['paragraphs']}  "
             f"tabelle={total['tables']}  "
             f"parametri={total['parameters']}  "
             f"tempo={total['duration_ms']}ms")

    # Costruisce FTS dopo estrazione (sempre, a meno di dry-run)
    if conn and not args.dry_run:
        build_fts(conn, log)
        conn.close()


if __name__ == "__main__":
    main()

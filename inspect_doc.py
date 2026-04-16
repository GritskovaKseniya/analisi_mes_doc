"""
inspect_doc.py — Step 0: ispezione struttura reale di un documento .docx

Uso:
    python inspect_doc.py <path_al_file.docx> [opzioni]

Opzioni:
    --limit N       mostra solo i primi N elementi (default: 100)
    --style NOME    filtra solo i paragrafi con questo stile Word
    --search TESTO  mostra solo paragrafi che contengono TESTO
    --tables        mostra anche il contenuto delle tabelle
    --stats         mostra solo le statistiche, non i paragrafi
    --params        mostra solo i parametri rilevati (con confidenza)
    --full          mostra testo completo (default: troncato a 120 caratteri)

Esempi:
    python inspect_doc.py "//rete-ud-2/.../Fontana_standard.r01.docx" --limit 50
    python inspect_doc.py "//rete-ud-2/.../MES_WEB_SERVIZI.r06.docx" --stats
    python inspect_doc.py "//rete-ud-2/.../MES_WEB_SERVIZI.r06.docx" --style "Heading 1"
    python inspect_doc.py "//rete-ud-2/.../MES_WEB_SERVIZI.r06.docx" --search "WMS_ABIL"
    python inspect_doc.py "//rete-ud-2/.../MES_WEB_SERVIZI.r06.docx" --params --limit 50
    python inspect_doc.py "//rete-ud-2/.../Fontana_standard.r01.docx" --params
"""

import argparse
import io
import re
import sys
from collections import Counter
from dataclasses import dataclass

# Forza UTF-8 sull'output — necessario su Windows con terminale cp1252
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

try:
    import docx
except ImportError:
    print("Errore: python-docx non installato. Esegui: pip install python-docx")
    sys.exit(1)


# ---------------------------------------------------------------------------
# Pattern di rilevamento parametri
# Ordine importante: il primo match vince. Mettere i più precisi prima.
# ---------------------------------------------------------------------------
PARAM_PATTERNS = [
    # 1. Shell export — il più specifico, quasi zero falsi positivi
    ("shell_export",
     re.compile(r'^export\s+([A-Z][A-Z0-9_]{2,})\s*=\s*(.+)$'),
     "linea_intera"),

    # 2. NOME_MAIUSCOLO=valore come intera linea (caso pulito)
    ("key_value_uppercase",
     re.compile(r'^([A-Z][A-Z0-9_]{2,})\s*=\s*(.+)$'),
     "linea_intera"),

    # 3. snake_case minuscolo = numero o stringa quoted (stile campi DB — risolve Fontana)
    # Valore vincolato: evita "se campo = descrizione"
    ("key_value_snakecase",
     re.compile(r'^([a-z][a-z0-9_]{2,})\s*=\s*([0-9]+|\'[^\']*\'|"[^"]*"|[A-Z])\s*$'),
     "linea_intera"),

    # 4. NOME_MAIUSCOLO=numero inline nel testo — cattura anche dentro frasi
    # Valore solo numerico: esclude NOME=testo narrativo
    ("inline_uppercase_param",
     re.compile(r'([A-Z][A-Z0-9_]{2,})\s*=\s*([0-9]+)'),
     "inline"),
]

# Parole iniziali che indicano contesto narrativo — abbassano la confidenza
NARRATIVE_PREFIXES = re.compile(
    r'^\s*(se |quando |qualora |esempio:|ad esempio|nel caso|if |nota:|nbb |nbb:|attenzione)',
    re.IGNORECASE
)

# Parole nel testo che segnalano che l'= è comparativo, non assegnativo
NARRATIVE_OPERATORS = re.compile(
    r'\b(è|sono|deve|vale|vuol dire|significa|ovvero|cioè|corrisponde)\b',
    re.IGNORECASE
)


@dataclass
class ParamMatch:
    pattern_name: str
    key: str
    value: str
    confidence: str          # "HIGH", "MEDIUM", "LOW"
    confidence_reason: str
    raw_text: str


def compute_confidence(text: str, pattern_name: str, key: str, value: str) -> tuple[str, str]:
    """
    Calcola la confidenza che questo match sia davvero un parametro di config.
    Ritorna (livello, motivo).

    Euristica (dalla nostra analisi):
    - Linea corta + valore numerico + nome con underscore → HIGH
    - Prefisso narrativo → scala di un livello
    - Operatori narrativi nel testo → LOW
    - inline_uppercase su linea corta → MEDIUM
    - inline_uppercase su linea lunga con narrativo → LOW
    """
    text_stripped = text.strip()
    reasons = []
    score = 0

    # Fattori positivi
    if "_" in key:
        score += 1
        reasons.append("nome_con_underscore")
    if re.match(r'^[0-9]+$', value):
        score += 1
        reasons.append("valore_numerico")
    if len(text_stripped) < 80:
        score += 1
        reasons.append("linea_corta")
    if pattern_name in ("shell_export", "key_value_uppercase", "key_value_snakecase"):
        score += 1
        reasons.append("pattern_linea_intera")

    # Fattori negativi
    if NARRATIVE_PREFIXES.match(text_stripped):
        score -= 2
        reasons.append("prefisso_narrativo")
    if NARRATIVE_OPERATORS.search(text_stripped):
        score -= 1
        reasons.append("operatore_narrativo")
    if pattern_name == "inline_uppercase_param" and len(text_stripped) > 100:
        score -= 1
        reasons.append("linea_lunga_inline")

    if score >= 3:
        level = "HIGH"
    elif score >= 1:
        level = "MEDIUM"
    else:
        level = "LOW"

    return level, "+".join(reasons) if reasons else "nessuno"


def detect_params(text: str) -> list[ParamMatch]:
    """Trova tutti i parametri nel testo. Può trovarne più di uno (caso inline)."""
    text_stripped = text.strip()
    results = []
    seen_keys = set()

    for name, pattern, match_type in PARAM_PATTERNS:
        if match_type == "linea_intera":
            m = pattern.match(text_stripped)
            if m:
                key, value = m.group(1), m.group(2).strip()
                if key not in seen_keys:
                    seen_keys.add(key)
                    conf, reason = compute_confidence(text_stripped, name, key, value)
                    results.append(ParamMatch(name, key, value, conf, reason, text_stripped))
                break  # linea intera: un solo match per linea, prendi il primo
        else:  # inline
            for m in pattern.finditer(text_stripped):
                key, value = m.group(1), m.group(2).strip()
                if key not in seen_keys:
                    seen_keys.add(key)
                    conf, reason = compute_confidence(text_stripped, name, key, value)
                    results.append(ParamMatch(name, key, value, conf, reason, text_stripped))

    return results


def truncate(text: str, width: int) -> str:
    text = text.replace('\n', ' ').replace('\r', '')
    return text if len(text) <= width else text[:width] + '…'


def conf_color(level: str) -> str:
    return {"HIGH": "\033[0;32m", "MEDIUM": "\033[0;33m", "LOW": "\033[0;31m"}.get(level, "")


RESET = "\033[0m"


def inspect(filepath: str, limit: int, style_filter: str | None,
            search: str | None, show_tables: bool, stats_only: bool,
            params_only: bool, full: bool):

    print(f"\n{'='*70}")
    print(f"  FILE: {filepath}")
    print(f"{'='*70}\n")

    try:
        doc = docx.Document(filepath)
    except Exception as e:
        print(f"Errore apertura file: {e}")
        sys.exit(1)

    # ------------------------------------------------------------------ stats
    style_counter: Counter = Counter()
    param_high = []
    param_medium = []
    param_low = []
    empty_count = 0
    short_count = 0
    heading_texts = []
    pattern_counter: Counter = Counter()

    for p in doc.paragraphs:
        style_counter[p.style.name] += 1
        text = p.text.strip()
        if not text:
            empty_count += 1
        elif len(text) <= 5:
            short_count += 1
        if p.style.name.startswith("Heading") or p.style.name.startswith("Titolo"):
            heading_texts.append((p.style.name, text))

        for pm in detect_params(text):
            pattern_counter[pm.pattern_name] += 1
            if pm.confidence == "HIGH":
                param_high.append(pm)
            elif pm.confidence == "MEDIUM":
                param_medium.append(pm)
            else:
                param_low.append(pm)

    total_params = len(param_high) + len(param_medium) + len(param_low)

    print(f"  Paragrafi totali : {len(doc.paragraphs)}")
    print(f"  Tabelle totali   : {len(doc.tables)}")
    print(f"  Paragrafi vuoti  : {empty_count}")
    print(f"  Paragrafi <=5ch  : {short_count}")
    print(f"  Parametri totali : {total_params}  "
          f"[\033[0;32mHIGH={len(param_high)}\033[0m  "
          f"\033[0;33mMEDIUM={len(param_medium)}\033[0m  "
          f"\033[0;31mLOW={len(param_low)}\033[0m]")
    print(f"  Per pattern      : {dict(pattern_counter)}")
    print()

    print("  Stili presenti (ordinati per frequenza):")
    for style, count in style_counter.most_common():
        bar = '█' * min(count // 10, 40)
        print(f"    {count:>6}  {style:<35} {bar}")
    print()

    if heading_texts:
        print(f"  Titoli Heading trovati ({len(heading_texts)}):")
        for style, text in heading_texts[:30]:
            indent = "  " * (int(style[-1]) - 1) if style[-1].isdigit() else ""
            print(f"    {indent}[{style}] {truncate(text, 80)}")
        if len(heading_texts) > 30:
            print(f"    ... altri {len(heading_texts) - 30} heading (usa --style 'Heading 1')")
    print()

    if stats_only:
        return

    # ------------------------------------------------------ modalità --params
    if params_only:
        print(f"{'─'*70}")
        print(f"  PARAMETRI RILEVATI (limit={limit})")
        print(f"  [verde=HIGH confidenza, giallo=MEDIUM, rosso=LOW]")
        print(f"{'─'*70}\n")

        all_params = param_high + param_medium + param_low
        shown = 0
        for pm in all_params:
            if shown >= limit:
                print(f"\n  ... altri {len(all_params) - shown} parametri. Aumenta --limit.\n")
                break
            color = conf_color(pm.confidence)
            key_padded = pm.key[:35].ljust(35)
            val_padded = pm.value[:20].ljust(20)
            pattern_short = pm.pattern_name[:22].ljust(22)
            print(f"  {color}{pm.confidence:<6}{RESET}  {key_padded}  = {val_padded}  "
                  f"[{pattern_short}]  {truncate(pm.raw_text, 60)}")
            shown += 1

        print(f"\n  RIEPILOGO PRECISION STIMATA:")
        print(f"  HIGH   {len(param_high):>4} — molto probabilmente veri parametri")
        print(f"  MEDIUM {len(param_medium):>4} — verificare manualmente una sample")
        print(f"  LOW    {len(param_low):>4} — probabilmente falsi positivi, da escludere")
        if total_params > 0:
            pct = len(param_high) / total_params * 100
            print(f"\n  Se usi solo HIGH: precision attesa >90% ({pct:.0f}% del totale)")
        return

    # -------------------------------------------------------------- paragrafi
    print(f"{'─'*70}")
    print(f"  PARAGRAFI (limit={limit}"
          + (f", stile='{style_filter}'" if style_filter else "")
          + (f", cerca='{search}'" if search else "")
          + ")")
    print(f"{'─'*70}\n")

    shown = 0
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        style = p.style.name

        if style_filter and style_filter.lower() not in style.lower():
            continue
        if search and search.lower() not in text.lower():
            continue
        if not text:
            continue

        if style.startswith("Heading") or style.startswith("Titolo"):
            level = style[-1] if style[-1].isdigit() else "?"
            prefix = f"\033[1;36m[H{level}]\033[0m"
        elif style == "List Paragraph":
            prefix = f"\033[0;33m[LIST]\033[0m"
        else:
            prefix = f"\033[0;37m[{style[:6]}]\033[0m"

        display_text = text if full else truncate(text, 120)

        params = detect_params(text)
        if params:
            best = params[0]
            color = conf_color(best.confidence)
            suffix = f"  {color}← PARAM {best.confidence} ({best.pattern_name}){RESET}"
        else:
            suffix = ""

        print(f"  {i:>5}  {prefix}  {display_text}{suffix}")
        shown += 1

        if shown >= limit:
            print(f"\n  ... aumenta --limit per vedere di più.\n")
            break

    # ----------------------------------------------------------------- tabelle
    if show_tables:
        print(f"\n{'─'*70}")
        print(f"  TABELLE (prime {min(limit, len(doc.tables))})")
        print(f"{'─'*70}\n")

        for t_idx, table in enumerate(doc.tables[:limit]):
            print(f"  Tabella {t_idx} — {len(table.rows)} righe × {len(table.columns)} colonne")
            for r_idx, row in enumerate(table.rows[:8]):
                cells = [truncate(cell.text.strip(), 35) for cell in row.cells]
                if cells and re.match(r'^[A-Z][A-Z0-9_]{2,}$', cells[0]):
                    marker = "\033[0;32m★\033[0m"
                elif cells and re.match(r'^[a-z][a-z0-9_]{2,}$', cells[0]):
                    marker = "\033[0;33m◆\033[0m"
                else:
                    marker = " "
                print(f"  {marker} r{r_idx}: {' | '.join(cells)}")
            if len(table.rows) > 8:
                print(f"    ... altre {len(table.rows) - 8} righe")
            print()


def main():
    parser = argparse.ArgumentParser(
        description="Ispeziona struttura di un documento .docx")
    parser.add_argument("filepath", help="Path al file .docx")
    parser.add_argument("--limit",  type=int, default=100,
                        help="Max elementi da mostrare (default: 100)")
    parser.add_argument("--style",  default=None,
                        help="Filtra per nome stile Word (es. 'Heading 1')")
    parser.add_argument("--search", default=None,
                        help="Filtra paragrafi contenenti questo testo")
    parser.add_argument("--tables", action="store_true",
                        help="Mostra anche contenuto tabelle")
    parser.add_argument("--stats",  action="store_true",
                        help="Solo statistiche, nessun dettaglio")
    parser.add_argument("--params", action="store_true",
                        help="Mostra solo i parametri rilevati con confidenza")
    parser.add_argument("--full",   action="store_true",
                        help="Testo completo (non troncato)")

    args = parser.parse_args()

    inspect(
        filepath=args.filepath,
        limit=args.limit,
        style_filter=args.style,
        search=args.search,
        show_tables=args.tables,
        stats_only=args.stats,
        params_only=args.params,
        full=args.full,
    )


if __name__ == "__main__":
    main()
